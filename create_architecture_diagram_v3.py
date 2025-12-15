#!/usr/bin/env python3
"""
Create a professional architecture diagram with properly oriented boxes and arrows
"""

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle
import matplotlib.patheffects as path_effects
from matplotlib.patches import Circle, Polygon
import numpy as np

# Set up the figure with high DPI
fig = plt.figure(figsize=(16, 12), dpi=300)
ax = fig.add_subplot(111)

# Set background color
fig.patch.set_facecolor("#FFFFFF")
ax.set_facecolor("#F8F9FA")

# Remove axes
ax.set_xlim(0, 16)
ax.set_ylim(0, 12)
ax.axis("off")

# Define professional color palette
colors = {
    "event": "#FF6B6B",  # Red
    "eventbridge": "#FF8C42",  # Orange
    "s3": "#4ECDC4",  # Teal
    "glue": "#FFE66D",  # Gold
    "bronze": "#B8860B",  # Dark goldenrod
    "silver": "#C0C0C0",  # Silver
    "gold": "#FFD700",  # Gold
    "athena": "#FF9900",  # AWS Orange
    "arrow": "#2C3E50",  # Dark blue-gray
    "text_dark": "#2C3E50",
    "text_light": "#ECF0F1",
}


def create_box(
    ax, x, y, width, height, text, color, subtext="", text_size=11, bold=True
):
    """Create a professional box with text"""
    # Main box
    box = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.08",
        edgecolor=color,
        facecolor=color,
        alpha=0.15,
        linewidth=2.5,
    )
    ax.add_patch(box)

    # Add border effect
    border_box = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.08",
        edgecolor=color,
        facecolor="none",
        linewidth=2.5,
        linestyle="-",
    )
    ax.add_patch(border_box)

    # Add main text
    weight = "bold" if bold else "normal"
    ax.text(
        x + width / 2,
        y + height / 2 + 0.15,
        text,
        ha="center",
        va="center",
        fontsize=text_size,
        fontweight=weight,
        color=colors["text_dark"],
    )

    # Add subtext
    if subtext:
        ax.text(
            x + width / 2,
            y + height / 2 - 0.25,
            subtext,
            ha="center",
            va="center",
            fontsize=8,
            style="italic",
            color="#555555",
        )

    return box


def create_arrow(ax, x1, y1, x2, y2, label="", curve=0, style="->"):
    """Create a styled arrow between components"""
    arrow = FancyArrowPatch(
        (x1, y1),
        (x2, y2),
        arrowstyle=style,
        mutation_scale=25,
        linewidth=2.5,
        color=colors["arrow"],
        connectionstyle=f"arc3,rad={curve}",
    )
    ax.add_patch(arrow)

    if label:
        mid_x, mid_y = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(
            mid_x + 0.15,
            mid_y + 0.25,
            label,
            fontsize=8,
            color=colors["arrow"],
            fontweight="bold",
            bbox=dict(
                boxstyle="round,pad=0.25",
                facecolor="white",
                alpha=0.9,
                edgecolor="none",
            ),
        )


def add_title(ax, title, subtitle=""):
    """Add title and subtitle"""
    title_text = ax.text(
        8,
        11.5,
        title,
        ha="center",
        va="top",
        fontsize=24,
        fontweight="bold",
        color=colors["text_dark"],
    )

    if subtitle:
        ax.text(
            8,
            11.0,
            subtitle,
            ha="center",
            va="top",
            fontsize=12,
            style="italic",
            color="#555555",
        )


# ============ TITLE ============
add_title(
    ax,
    "Event Analytics Pipeline Architecture",
    "Bronze → Silver → Gold Medallion Pattern",
)

# ============ LAYER 1: EVENT GENERATION (TOP) ============
y_top = 10

# Lambda
create_box(
    ax, 1, y_top, 2, 0.7, "AWS Lambda", colors["event"], "Event Generator", text_size=10
)

# Arrow from Lambda to EventBridge
create_arrow(ax, 3, y_top + 0.35, 3.8, y_top + 0.35, style="->")

# EventBridge
create_box(
    ax,
    3.8,
    y_top,
    2,
    0.7,
    "EventBridge",
    colors["eventbridge"],
    "Scheduler (5 min)",
    text_size=10,
)

# Arrow from EventBridge to S3
create_arrow(ax, 5.8, y_top + 0.35, 6.8, y_top + 0.35, style="->")

# Source S3 Bucket
create_box(
    ax,
    6.8,
    y_top,
    2.5,
    0.7,
    "Source S3",
    colors["s3"],
    "500K-750K events",
    text_size=10,
)

# Data volume annotation
vol_box = Rectangle(
    (10, y_top), 4, 0.7, facecolor="#FFF3CD", edgecolor="#FF9800", linewidth=2
)
ax.add_patch(vol_box)
ax.text(
    12,
    y_top + 0.55,
    "Data Volume",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#FF6F00",
)
ax.text(
    12,
    y_top + 0.15,
    "~50 GB per day",
    fontsize=10,
    fontweight="bold",
    ha="center",
    color="#FF6F00",
)

# ============ ARROWS DOWNWARD ============
# From S3 to Glue
create_arrow(ax, 8.05, y_top, 6, 8.5, curve=0, style="->")

# ============ LAYER 2: ETL PROCESSING (MIDDLE) ============
y_glue = 7.5

# Glue Job - Large prominent box
glue_box = FancyBboxPatch(
    (3.5, y_glue),
    5,
    1,
    boxstyle="round,pad=0.1",
    edgecolor=colors["glue"],
    facecolor=colors["glue"],
    alpha=0.2,
    linewidth=3,
)
ax.add_patch(glue_box)

ax.text(
    6,
    y_glue + 0.75,
    "AWS Glue ETL Job",
    fontsize=12,
    fontweight="bold",
    ha="center",
    color=colors["text_dark"],
)
ax.text(
    6,
    y_glue + 0.35,
    "Job Bookmarks Enabled | Incremental Processing",
    fontsize=9,
    ha="center",
    color="#1565C0",
    fontweight="bold",
)

# Performance badge on right
perf_box = Rectangle(
    (10, y_glue + 0.15), 4, 0.7, facecolor="#E8F5E9", edgecolor="#4CAF50", linewidth=2
)
ax.add_patch(perf_box)
ax.text(
    12,
    y_glue + 0.65,
    "⚡ 20% Faster",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#2E7D32",
)
ax.text(12, y_glue + 0.25, "Incremental Runs", fontsize=8, ha="center", color="#1B5E20")

# ============ ARROWS DOWN FROM GLUE TO MEDALLION ============
# From Glue to Bronze
create_arrow(ax, 4.5, y_glue, 2.2, 6.2, curve=-0.1, style="->")

# From Glue to Silver (straight down)
create_arrow(ax, 6, y_glue, 6, 6.2, curve=0, style="->", label="Transform")

# From Glue to Gold
create_arrow(ax, 7.5, y_glue, 9.8, 6.2, curve=0.1, style="->")

# ============ LAYER 3: MEDALLION ARCHITECTURE ============
y_medallion = 5.2

# Bronze Layer (LEFT)
create_box(
    ax,
    0.5,
    y_medallion,
    3,
    1,
    "BRONZE LAYER\nRaw Parquet\nAppend Mode",
    colors["bronze"],
    "Raw JSON → Parquet",
    text_size=10,
    bold=True,
)

# Silver Layer (CENTER)
create_box(
    ax,
    4.5,
    y_medallion,
    3,
    1,
    "SILVER LAYER\nCleaned & Deduplicated\nOverwrite Mode",
    colors["silver"],
    "Parse • Transform • Dedupe",
    text_size=10,
    bold=True,
)

# Gold Layer (RIGHT)
create_box(
    ax,
    8.5,
    y_medallion,
    3,
    1,
    "GOLD LAYER\nAnalytics Ready\n5 Tables",
    colors["gold"],
    "Pre-aggregated",
    text_size=10,
    bold=True,
)

# ============ ARROWS BETWEEN MEDALLION LAYERS ============
# Bronze to Silver
create_arrow(ax, 3.5, y_medallion + 0.5, 4.5, y_medallion + 0.5, style="->")

# Silver to Gold
create_arrow(ax, 7.5, y_medallion + 0.5, 8.5, y_medallion + 0.5, style="->")

# ============ LAYER 4: GOLD TABLES ============
y_tables = 3.5

table_names = [
    "hourly_\nrevenue",
    "top_\nproducts",
    "category_\nperf",
    "user_\nactivity",
    "conversion_\nfunnel",
]
table_x_positions = [1.2, 3.0, 4.8, 6.6, 8.4]

for table_name, x_pos in zip(table_names, table_x_positions):
    table_box = FancyBboxPatch(
        (x_pos, y_tables),
        1.4,
        0.8,
        boxstyle="round,pad=0.06",
        edgecolor=colors["gold"],
        facecolor=colors["gold"],
        alpha=0.25,
        linewidth=2,
    )
    ax.add_patch(table_box)
    ax.text(
        x_pos + 0.7,
        y_tables + 0.4,
        table_name,
        ha="center",
        va="center",
        fontsize=8,
        fontweight="bold",
        color=colors["text_dark"],
    )

# Arrow from Gold to Tables
create_arrow(ax, 10, y_medallion, 6, y_tables + 0.8, curve=-0.15, style="->")
ax.text(8.5, 4.3, "Aggregations", fontsize=8, color="#555555", style="italic")

# ============ KPI BADGES (RIGHT SIDE) ============
y_badges = y_medallion

# Freshness badge
fresh_box = Rectangle(
    (11, y_badges + 0.8), 4, 0.7, facecolor="#E3F2FD", edgecolor="#2196F3", linewidth=2
)
ax.add_patch(fresh_box)
ax.text(
    13,
    y_badges + 1.35,
    "Data Freshness",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#0D47A1",
)
ax.text(
    13,
    y_badges + 0.95,
    "5-30 minutes",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#1565C0",
)

# Cost badge
cost_box = Rectangle(
    (11, y_badges), 4, 0.7, facecolor="#F3E5F5", edgecolor="#9C27B0", linewidth=2
)
ax.add_patch(cost_box)
ax.text(
    13,
    y_badges + 0.55,
    "Cost Efficient",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#6A1B9A",
)
ax.text(
    13,
    y_badges + 0.15,
    "~$47/month",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#7B1FA2",
)

# ============ ARROW DOWN TO ATHENA ============
create_arrow(ax, 6, y_tables, 6, 1.8, curve=0, style="->", label="Query")

# ============ LAYER 5: QUERY ENGINE (BOTTOM) ============
y_athena = 0.8

athena_box = FancyBboxPatch(
    (2, y_athena),
    8,
    0.9,
    boxstyle="round,pad=0.1",
    edgecolor=colors["athena"],
    facecolor=colors["athena"],
    alpha=0.2,
    linewidth=3,
)
ax.add_patch(athena_box)

ax.text(
    6,
    y_athena + 0.65,
    "Amazon Athena - SQL Query Engine",
    fontsize=12,
    fontweight="bold",
    ha="center",
    color=colors["text_dark"],
)
ax.text(
    6,
    y_athena + 0.25,
    "Real-time Analytics on Gold Tables",
    fontsize=9,
    ha="center",
    color="#555555",
    style="italic",
)

# Query performance badge
query_box = Rectangle(
    (11, y_athena + 0.1), 4, 0.7, facecolor="#F1F8E9", edgecolor="#7CB342", linewidth=2
)
ax.add_patch(query_box)
ax.text(
    13,
    y_athena + 0.65,
    "Query Speed",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#558B2F",
)
ax.text(
    13,
    y_athena + 0.25,
    "<1 second",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#689F38",
)

# ============ BOTTOM FLOW DESCRIPTION ============
flow_y = -0.2
ax.text(
    8,
    flow_y - 0.3,
    "Data Flow: Events (Lambda/S3) → ETL (Glue Bookmarks) → Bronze (Raw) → Silver (Clean) → Gold (Analytics) → Athena (Query)",
    ha="center",
    va="top",
    fontsize=9,
    style="italic",
    color="#333333",
    fontweight="bold",
    bbox=dict(
        boxstyle="round,pad=0.6",
        facecolor="#ECEFF1",
        alpha=0.9,
        edgecolor="#455A64",
        linewidth=1.5,
    ),
)

plt.tight_layout()
plt.savefig(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    edgecolor="none",
)
print("✅ Improved architecture diagram created: architecture_diagram.png")
plt.close()

# Now update the Word document with the diagram
print("\n📝 Updating Word document with improved diagram...")

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing Word document
doc = Document(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)

# Find and remove old diagram placeholder paragraphs
paragraphs_to_remove = []
for i, para in enumerate(doc.paragraphs):
    if (
        "Architecture diagram" in para.text
        or "[Note:" in para.text
        or para.text.strip() == ""
    ):
        # Check if next paragraph has an image
        if i < len(doc.paragraphs) - 1:
            for run in doc.paragraphs[i + 1].runs:
                if "architecture_diagram" in str(run):
                    paragraphs_to_remove.append(i + 1)

# Remove old paragraphs (in reverse order to maintain indices)
for i in reversed(paragraphs_to_remove):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

# Find insertion point (after "Solution Architecture" heading)
insert_index = None
for i, para in enumerate(doc.paragraphs):
    if "Solution Architecture" in para.text:
        insert_index = i + 1
        break

# Add new paragraph with improved image
if insert_index:
    # Insert paragraph at correct location
    insertion_para = doc.paragraphs[insert_index]._element
    new_para_element = doc.add_paragraph()._element
    insertion_para.addprevious(new_para_element)
    new_para = doc.paragraphs[insert_index]
else:
    new_para = doc.add_paragraph()

new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = new_para.add_run()
run.add_picture(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
    width=Inches(6.5),
)

# Save updated document
doc.save(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)
print("✅ Word document updated with improved diagram")

print("\n" + "=" * 80)
print("IMPROVED ARCHITECTURE DIAGRAM COMPLETE")
print("=" * 80)
print("\n🎯 Improvements Made:")
print("  ✓ Better left-to-right flow")
print("  ✓ Proper vertical alignment")
print("  ✓ Clear downward data flow")
print("  ✓ Medallion layers properly positioned")
print("  ✓ Arrows oriented correctly")
print("  ✓ All boxes clearly aligned")
print("  ✓ Professional layout")
print("\n✅ Ready for submission!")
