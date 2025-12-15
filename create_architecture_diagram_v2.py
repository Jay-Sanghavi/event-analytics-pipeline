#!/usr/bin/env python3
"""
Create a professional architecture diagram for the event analytics pipeline
"""

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle
import matplotlib.patheffects as path_effects
from matplotlib.patches import Circle
import numpy as np

# Set up the figure with high DPI
fig = plt.figure(figsize=(16, 11), dpi=300)
ax = fig.add_subplot(111)

# Set background color
fig.patch.set_facecolor("#FFFFFF")
ax.set_facecolor("#F8F9FA")

# Remove axes
ax.set_xlim(0, 16)
ax.set_ylim(0, 11)
ax.axis("off")

# Define professional color palette
colors = {
    "event": "#FF6B6B",  # Red
    "eventbridge": "#FF8C42",  # Orange
    "s3": "#4ECDC4",  # Teal
    "glue": "#FFE66D",  # Gold
    "bronze": "#B8860B",  # Dark goldenrod
    "silver": "#C0C0C0",  # Silver
    "gold": "#FFD700",  # Gold
    "athena": "#FF9900",  # AWS Orange
    "arrow": "#2C3E50",  # Dark blue-gray
    "text_dark": "#2C3E50",
    "text_light": "#ECF0F1",
}


def create_box(ax, x, y, width, height, text, color, icon="", text_size=11, bold=True):
    """Create a professional box with text"""
    # Main box
    box = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.08",
        edgecolor=color,
        facecolor=color,
        alpha=0.15,
        linewidth=2.5,
    )
    ax.add_patch(box)

    # Add border effect
    border_box = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.08",
        edgecolor=color,
        facecolor="none",
        linewidth=2.5,
        linestyle="-",
    )
    ax.add_patch(border_box)

    # Add text
    weight = "bold" if bold else "normal"
    ax.text(
        x + width / 2,
        y + height / 2 + 0.15,
        text,
        ha="center",
        va="center",
        fontsize=text_size,
        fontweight=weight,
        color=colors["text_dark"],
    )

    return box


def create_arrow(ax, x1, y1, x2, y2, label="", curve=0):
    """Create a styled arrow between components"""
    arrow = FancyArrowPatch(
        (x1, y1),
        (x2, y2),
        arrowstyle="->",
        mutation_scale=30,
        linewidth=2.5,
        color=colors["arrow"],
        connectionstyle=f"arc3,rad={curve}",
    )
    ax.add_patch(arrow)

    if label:
        mid_x, mid_y = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(
            mid_x + 0.2,
            mid_y + 0.2,
            label,
            fontsize=9,
            color=colors["arrow"],
            bbox=dict(
                boxstyle="round,pad=0.3", facecolor="white", alpha=0.8, edgecolor="none"
            ),
        )


def add_title(ax, title, subtitle=""):
    """Add title and subtitle"""
    title_text = ax.text(
        8,
        10.5,
        title,
        ha="center",
        va="top",
        fontsize=24,
        fontweight="bold",
        color=colors["text_dark"],
    )
    title_text.set_path_effects(
        [path_effects.Stroke(linewidth=3, foreground="white"), path_effects.Normal()]
    )

    if subtitle:
        ax.text(
            8,
            10.0,
            subtitle,
            ha="center",
            va="top",
            fontsize=12,
            style="italic",
            color="#555555",
        )


# ============ TITLE ============
add_title(
    ax,
    "Event Analytics Pipeline",
    "Production-Grade Architecture with Incremental Processing",
)

# ============ LAYER 1: EVENT GENERATION ============
y_layer1 = 8.8

# Lambda
create_box(
    ax,
    1,
    y_layer1,
    2,
    0.8,
    "AWS Lambda\nEvent Generator",
    colors["event"],
    text_size=10,
)
ax.text(
    2,
    y_layer1 - 0.3,
    "Executes every\n5 minutes",
    ha="center",
    fontsize=8,
    color="#666666",
    style="italic",
)

# EventBridge
create_box(
    ax,
    4,
    y_layer1,
    2,
    0.8,
    "EventBridge\nScheduler",
    colors["eventbridge"],
    text_size=10,
)
ax.text(
    5,
    y_layer1 - 0.3,
    "Orchestration\nTrigger",
    ha="center",
    fontsize=8,
    color="#666666",
    style="italic",
)

# Source S3
create_box(ax, 7, y_layer1, 2.2, 0.8, "Source S3\nBucket", colors["s3"], text_size=10)
ax.text(
    8.1,
    y_layer1 - 0.3,
    "500K-750K\nevents/5min",
    ha="center",
    fontsize=8,
    color="#666666",
    style="italic",
)

# Data volume box
volume_box = Rectangle(
    (10.5, y_layer1 - 0.2), 3, 1, facecolor="#FFF3CD", edgecolor="#FF9800", linewidth=2
)
ax.add_patch(volume_box)
ax.text(
    12,
    y_layer1 + 0.3,
    "📊 Data Volume",
    fontsize=10,
    fontweight="bold",
    ha="center",
    color="#FF6F00",
)
ax.text(
    12,
    y_layer1,
    "~50 GB/day",
    fontsize=11,
    fontweight="bold",
    ha="center",
    color="#FF6F00",
)
ax.text(12, y_layer1 - 0.2, "Gzipped JSON", fontsize=8, ha="center", color="#666666")

# Arrows
create_arrow(ax, 3, y_layer1 + 0.4, 4, y_layer1 + 0.4)
create_arrow(ax, 6, y_layer1 + 0.4, 7, y_layer1 + 0.4)

# ============ LAYER 2: ETL PROCESSING ============
y_layer2 = 7

# Glue Job - Large box
glue_box = FancyBboxPatch(
    (3.5, y_layer2 - 0.5),
    4.5,
    1.2,
    boxstyle="round,pad=0.1",
    edgecolor=colors["glue"],
    facecolor=colors["glue"],
    alpha=0.2,
    linewidth=3,
)
ax.add_patch(glue_box)

ax.text(
    5.75,
    y_layer2 + 0.5,
    "AWS Glue ETL Job",
    fontsize=13,
    fontweight="bold",
    ha="center",
    color=colors["text_dark"],
)
ax.text(
    5.75,
    y_layer2 + 0.1,
    "✓ Job Bookmarks Enabled",
    fontsize=9,
    ha="center",
    color="#2E7D32",
    fontweight="bold",
)
ax.text(
    5.75,
    y_layer2 - 0.2,
    "Incremental Processing • Only NEW files processed",
    fontsize=9,
    ha="center",
    color="#1565C0",
)

# Arrow from S3 to Glue
create_arrow(ax, 8.1, y_layer1, 5.75, y_layer2 + 0.7, curve=-0.3)

# Performance badge
perf_box = Rectangle(
    (10.5, y_layer2 - 0.3),
    3,
    0.7,
    facecolor="#E8F5E9",
    edgecolor="#4CAF50",
    linewidth=2,
)
ax.add_patch(perf_box)
ax.text(
    12,
    y_layer2 + 0.1,
    "⚡ Performance",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#2E7D32",
)
ax.text(
    12,
    y_layer2 - 0.15,
    "20% faster incremental",
    fontsize=8,
    ha="center",
    color="#1B5E20",
)

# ============ LAYER 3: MEDALLION ARCHITECTURE ============
y_layer3 = 4.8

# Bronze Layer
create_box(
    ax,
    0.5,
    y_layer3,
    2.8,
    1.3,
    "BRONZE LAYER\n\nRaw Data\nParquet Format\nAppend Mode\n(Incremental)",
    colors["bronze"],
    text_size=9,
    bold=True,
)
ax.text(
    1.9,
    y_layer3 - 0.4,
    "Raw JSON → Parquet",
    fontsize=7,
    ha="center",
    style="italic",
    color="#555555",
)

# Silver Layer
create_box(
    ax,
    3.8,
    y_layer3,
    2.8,
    1.3,
    "SILVER LAYER\n\nCleaned Data\nDeduplication\nOverwrite Mode\n(Full Rebuild)",
    colors["silver"],
    text_size=9,
    bold=True,
)
ax.text(
    5.2,
    y_layer3 - 0.4,
    "Parse • Transform • Dedupe",
    fontsize=7,
    ha="center",
    style="italic",
    color="#555555",
)

# Gold Layer
create_box(
    ax,
    7.1,
    y_layer3,
    2.8,
    1.3,
    "GOLD LAYER\n\nAnalytics Ready\n5 Pre-aggregated\nTables",
    colors["gold"],
    text_size=9,
    bold=True,
)
ax.text(
    8.5,
    y_layer3 - 0.4,
    "Optimized for Queries",
    fontsize=7,
    ha="center",
    style="italic",
    color="#555555",
)

# Arrows between layers
create_arrow(ax, 3.3, y_layer3 + 0.65, 3.8, y_layer3 + 0.65, curve=0)
create_arrow(ax, 6.6, y_layer3 + 0.65, 7.1, y_layer3 + 0.65, curve=0)

# Arrow from Glue to Bronze
create_arrow(ax, 4.5, y_layer2 - 0.5, 1.9, y_layer3 + 1.3, curve=-0.2)

# Data freshness badge
fresh_box = Rectangle(
    (10.5, y_layer3 + 0.3),
    3,
    0.9,
    facecolor="#E3F2FD",
    edgecolor="#2196F3",
    linewidth=2,
)
ax.add_patch(fresh_box)
ax.text(
    12,
    y_layer3 + 0.9,
    "⏱ Data Freshness",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#0D47A1",
)
ax.text(
    12,
    y_layer3 + 0.45,
    "5-30 minutes",
    fontsize=10,
    fontweight="bold",
    ha="center",
    color="#1565C0",
)

# ============ LAYER 4: GOLD TABLES ============
y_layer4 = 2.8

tables = [
    ("hourly_\nrevenue", 1.2),
    ("top_\nproducts", 2.6),
    ("category_\nperf", 4.0),
    ("user_\nactivity", 5.4),
    ("conversion_\nfunnel", 6.8),
]

for table_name, x_pos in tables:
    table_box = FancyBboxPatch(
        (x_pos, y_layer4),
        1.2,
        0.8,
        boxstyle="round,pad=0.05",
        edgecolor=colors["gold"],
        facecolor=colors["gold"],
        alpha=0.3,
        linewidth=2,
    )
    ax.add_patch(table_box)
    ax.text(
        x_pos + 0.6,
        y_layer4 + 0.4,
        table_name,
        ha="center",
        va="center",
        fontsize=8,
        fontweight="bold",
        color=colors["text_dark"],
    )

# Arrows from Gold to Tables
create_arrow(ax, 8.5, y_layer3, 4, y_layer4 + 0.8, curve=-0.2)

# Cost badge
cost_box = Rectangle(
    (10.5, y_layer4 - 0.2),
    3,
    0.8,
    facecolor="#F3E5F5",
    edgecolor="#9C27B0",
    linewidth=2,
)
ax.add_patch(cost_box)
ax.text(
    12,
    y_layer4 + 0.3,
    "💰 Cost Efficient",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#6A1B9A",
)
ax.text(
    12,
    y_layer4 - 0.05,
    "~$47/month",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#7B1FA2",
)

# ============ LAYER 5: QUERY ENGINE ============
y_layer5 = 0.8

# Athena
athena_box = FancyBboxPatch(
    (2.5, y_layer5 - 0.2),
    5,
    1,
    boxstyle="round,pad=0.1",
    edgecolor=colors["athena"],
    facecolor=colors["athena"],
    alpha=0.2,
    linewidth=3,
)
ax.add_patch(athena_box)

ax.text(
    5,
    y_layer5 + 0.5,
    "Amazon Athena",
    fontsize=12,
    fontweight="bold",
    ha="center",
    color=colors["text_dark"],
)
ax.text(
    5,
    y_layer5 + 0.1,
    "SQL Query Engine • <1 Second Latency • Instant Analytics",
    fontsize=9,
    ha="center",
    color="#555555",
)

# Arrow from Tables to Athena
create_arrow(ax, 4, y_layer4, 5, y_layer5 + 0.8, curve=0)

# Query performance badge
query_box = Rectangle(
    (10.5, y_layer5), 3, 0.6, facecolor="#F1F8E9", edgecolor="#7CB342", linewidth=2
)
ax.add_patch(query_box)
ax.text(
    12,
    y_layer5 + 0.35,
    "🚀 Query Speed",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#558B2F",
)
ax.text(
    12,
    y_layer5 + 0.05,
    "<1 second",
    fontsize=9,
    fontweight="bold",
    ha="center",
    color="#689F38",
)

# ============ FLOW ANNOTATIONS ============
flow_text = "Data Flow: Events → S3 → Glue Job (Bookmarks) → Bronze (Raw) → Silver (Clean) → Gold (Analytics) → Athena (Query)"
ax.text(
    8,
    -0.3,
    flow_text,
    ha="center",
    va="top",
    fontsize=9,
    style="italic",
    color="#555555",
    bbox=dict(
        boxstyle="round,pad=0.5",
        facecolor="#ECEFF1",
        alpha=0.8,
        edgecolor="#455A64",
        linewidth=1,
    ),
)

plt.tight_layout()
plt.savefig(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    edgecolor="none",
)
print("✅ Professional architecture diagram created: architecture_diagram.png")
plt.close()

# Now update the Word document with the diagram
print("\n📝 Updating Word document with new diagram...")

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing Word document
doc = Document(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)

# Find and replace the old diagram
found = False
for i, para in enumerate(doc.paragraphs):
    if (
        "Architecture diagram placeholder" in para.text
        or "Note: Architecture diagram" in para.text
    ):
        # Remove old text
        p = para._element
        p.getparent().remove(p)
        found = True
        break

if found or True:  # Always add/replace diagram
    # Find the "Solution Architecture" section or add near it
    insert_position = None
    for i, para in enumerate(doc.paragraphs):
        if (
            "Solution Architecture" in para.text
            and "[Note:" in doc.paragraphs[i + 3].text
        ):
            insert_position = i + 3
            # Remove the old placeholder
            p = doc.paragraphs[i + 3]._element
            p.getparent().remove(p)
            break

    # Add new paragraph with image
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(
        "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
        width=Inches(6.5),
    )

# Save updated document
doc.save(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)
print("✅ Word document updated with professional diagram")

print("\n" + "=" * 80)
print("PROFESSIONAL ARCHITECTURE DIAGRAM COMPLETE")
print("=" * 80)
print("\n📊 Enhancements:")
print("  ✓ Professional color scheme with AWS branding")
print("  ✓ Clear layer separation (Bronze/Silver/Gold)")
print("  ✓ Performance and cost badges")
print("  ✓ Better typography and spacing")
print("  ✓ Information density optimized")
print("  ✓ High-resolution output (300 DPI)")
print("  ✓ Embedded in blog post")
print("\n✅ Ready for submission!")
