#!/usr/bin/env python3
"""
Create a visual architecture diagram for the event analytics pipeline
"""

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.lines as mlines

# Create figure and axis
fig, ax = plt.subplots(1, 1, figsize=(14, 10))
ax.set_xlim(0, 10)
ax.set_ylim(0, 10)
ax.axis("off")

# Define colors
color_event = "#FF6B6B"  # Red
color_s3 = "#4ECDC4"  # Teal
color_glue = "#FFE66D"  # Yellow
color_database = "#95E1D3"  # Light Teal
color_query = "#A8E6CF"  # Green
color_arrow = "#333333"  # Dark gray

# Title
ax.text(
    5,
    9.5,
    "Event Analytics Pipeline Architecture",
    ha="center",
    va="top",
    fontsize=18,
    fontweight="bold",
)
ax.text(
    5,
    9.1,
    "Bronze → Silver → Gold Medallion Architecture",
    ha="center",
    va="top",
    fontsize=12,
    style="italic",
    color="#555555",
)

# ============ ROW 1: EVENT GENERATION ============
# Lambda Event Generator
lambda_box = FancyBboxPatch(
    (0.5, 7.5),
    1.5,
    0.6,
    boxstyle="round,pad=0.1",
    edgecolor=color_event,
    facecolor=color_event,
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(lambda_box)
ax.text(
    1.25,
    7.8,
    "Lambda\nEvent Generator",
    ha="center",
    va="center",
    fontsize=10,
    fontweight="bold",
)

# EventBridge
bridge_box = FancyBboxPatch(
    (2.5, 7.5),
    1.5,
    0.6,
    boxstyle="round,pad=0.1",
    edgecolor="#FF8C42",
    facecolor="#FF8C42",
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(bridge_box)
ax.text(
    3.25,
    7.8,
    "EventBridge\nScheduler\n(5 min)",
    ha="center",
    va="center",
    fontsize=10,
    fontweight="bold",
)

# Source S3 Bucket
s3_source = FancyBboxPatch(
    (5.5, 7.5),
    1.8,
    0.6,
    boxstyle="round,pad=0.1",
    edgecolor=color_s3,
    facecolor=color_s3,
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(s3_source)
ax.text(
    6.4,
    7.8,
    "Source S3 Bucket\n(Raw JSON)",
    ha="center",
    va="center",
    fontsize=10,
    fontweight="bold",
)

# Arrows: Lambda -> EventBridge -> S3
arrow1 = FancyArrowPatch(
    (2, 7.8),
    (2.5, 7.8),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2,
    color=color_arrow,
)
ax.add_patch(arrow1)

arrow2 = FancyArrowPatch(
    (4, 7.8),
    (5.5, 7.8),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2,
    color=color_arrow,
)
ax.add_patch(arrow2)

# Event volume annotation
ax.text(
    6.4,
    7.1,
    "500K-750K events/5min\n(~50 GB/day)",
    ha="center",
    va="top",
    fontsize=9,
    style="italic",
    color="#555555",
)

# ============ ROW 2: GLUE JOB ============
glue_box = FancyBboxPatch(
    (2.5, 5.5),
    3,
    1.2,
    boxstyle="round,pad=0.1",
    edgecolor=color_glue,
    facecolor=color_glue,
    alpha=0.3,
    linewidth=3,
)
ax.add_patch(glue_box)
ax.text(
    4, 6.4, "AWS Glue ETL Job", ha="center", va="center", fontsize=12, fontweight="bold"
)
ax.text(
    4,
    6.05,
    "(Job Bookmarks Enabled)",
    ha="center",
    va="center",
    fontsize=9,
    style="italic",
)
ax.text(
    4,
    5.7,
    "Incremental Processing\nOnly NEW files processed",
    ha="center",
    va="center",
    fontsize=9,
    color="#2C5F2D",
)

# Arrow from S3 to Glue
arrow3 = FancyArrowPatch(
    (6.4, 7.5),
    (5.2, 6.7),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2.5,
    color=color_arrow,
)
ax.add_patch(arrow3)
ax.text(6.0, 7.0, "read", fontsize=9, color=color_arrow)

# ============ ROW 3: MEDALLION LAYERS ============
y_layer = 3.8

# Bronze Layer
bronze_box = FancyBboxPatch(
    (0.3, y_layer),
    2.5,
    1.2,
    boxstyle="round,pad=0.1",
    edgecolor="#C0756B",
    facecolor="#C0756B",
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(bronze_box)
ax.text(
    1.55,
    y_layer + 0.85,
    "Bronze Layer",
    ha="center",
    va="center",
    fontsize=11,
    fontweight="bold",
)
ax.text(
    1.55,
    y_layer + 0.4,
    "Raw Parquet\nAppend Mode\n(Incremental)",
    ha="center",
    va="center",
    fontsize=8,
)

# Arrow from Glue to Bronze
arrow_bronze = FancyArrowPatch(
    (3, 5.5),
    (1.8, 5.0),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2.5,
    color=color_arrow,
)
ax.add_patch(arrow_bronze)

# Silver Layer
silver_box = FancyBboxPatch(
    (3.75, y_layer),
    2.5,
    1.2,
    boxstyle="round,pad=0.1",
    edgecolor="#A8A8A8",
    facecolor="#A8A8A8",
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(silver_box)
ax.text(
    5,
    y_layer + 0.85,
    "Silver Layer",
    ha="center",
    va="center",
    fontsize=11,
    fontweight="bold",
)
ax.text(
    5,
    y_layer + 0.4,
    "Cleaned & Dedup\nOverwrite Mode\n(Rebuild)",
    ha="center",
    va="center",
    fontsize=8,
)

# Arrow from Bronze to Silver
arrow_silver = FancyArrowPatch(
    (2.8, 4.4),
    (3.75, 4.4),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2.5,
    color=color_arrow,
)
ax.add_patch(arrow_silver)

# Gold Layer
gold_box = FancyBboxPatch(
    (7.2, y_layer),
    2.5,
    1.2,
    boxstyle="round,pad=0.1",
    edgecolor="#DAA520",
    facecolor="#DAA520",
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(gold_box)
ax.text(
    8.45,
    y_layer + 0.85,
    "Gold Layer",
    ha="center",
    va="center",
    fontsize=11,
    fontweight="bold",
)
ax.text(
    8.45,
    y_layer + 0.4,
    "Pre-aggregated\n5 Analytics Tables\n(Query-ready)",
    ha="center",
    va="center",
    fontsize=8,
)

# Arrow from Silver to Gold
arrow_gold = FancyArrowPatch(
    (6.25, 4.4),
    (7.2, 4.4),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2.5,
    color=color_arrow,
)
ax.add_patch(arrow_gold)

# ============ ROW 4: GOLD TABLES ============
y_tables = 2.4

tables = [
    "hourly_revenue",
    "top_products",
    "category_perf",
    "user_activity",
    "conversion_funnel",
]

table_x_positions = [0.8, 2.2, 3.6, 5.0, 6.4]

for i, (table, x_pos) in enumerate(zip(tables, table_x_positions)):
    table_box = FancyBboxPatch(
        (x_pos, y_tables - 0.3),
        1.2,
        0.5,
        boxstyle="round,pad=0.05",
        edgecolor=color_query,
        facecolor=color_query,
        alpha=0.25,
        linewidth=1.5,
    )
    ax.add_patch(table_box)
    ax.text(
        x_pos + 0.6,
        y_tables,
        table.replace("_", "\n"),
        ha="center",
        va="center",
        fontsize=8,
        fontweight="bold",
    )

# Arrow from Gold to Tables
arrow_tables = FancyArrowPatch(
    (8.45, 3.8),
    (4, 2.7),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2,
    color=color_arrow,
    connectionstyle="arc3,rad=0.3",
)
ax.add_patch(arrow_tables)

# ============ ROW 5: QUERY LAYER ============
# Amazon Athena
athena_box = FancyBboxPatch(
    (2, 0.5),
    4,
    0.7,
    boxstyle="round,pad=0.1",
    edgecolor="#FF9900",
    facecolor="#FF9900",
    alpha=0.3,
    linewidth=2,
)
ax.add_patch(athena_box)
ax.text(
    4,
    0.9,
    "Amazon Athena (SQL Queries)\n<1 second query latency",
    ha="center",
    va="center",
    fontsize=11,
    fontweight="bold",
)

# Arrow from Tables to Athena
arrow_athena = FancyArrowPatch(
    (4, 2.1),
    (4, 1.2),
    arrowstyle="->",
    mutation_scale=20,
    linewidth=2.5,
    color=color_arrow,
)
ax.add_patch(arrow_athena)

# ============ ANNOTATIONS & LEGEND ============
# Data freshness annotation
ax.text(
    9.2,
    5.5,
    "Data Freshness:\n5-30 min",
    ha="left",
    va="top",
    fontsize=9,
    style="italic",
    bbox=dict(boxstyle="round", facecolor="#FFFFCC", alpha=0.7),
)

# Cost annotation
ax.text(
    9.2,
    4.3,
    "Monthly Cost:\n~$47",
    ha="left",
    va="top",
    fontsize=9,
    style="italic",
    bbox=dict(boxstyle="round", facecolor="#E8F5E9", alpha=0.7),
)

# Performance annotation
ax.text(
    9.2,
    3.1,
    "Performance:\n20% faster\nincremental runs",
    ha="left",
    va="top",
    fontsize=9,
    style="italic",
    bbox=dict(boxstyle="round", facecolor="#E3F2FD", alpha=0.7),
)

# Add processing flow legend
ax.text(
    0.3,
    0.1,
    "Processing Flow: Event → Raw Data → Transform (Glue) → Bronze (Raw) → Silver (Clean) → Gold (Analytics) → Query (Athena)",
    ha="left",
    va="bottom",
    fontsize=9,
    bbox=dict(boxstyle="round", facecolor="#F5F5F5", alpha=0.8),
)

plt.tight_layout()
plt.savefig(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    edgecolor="none",
)
print("✅ Architecture diagram created: architecture_diagram.png")
plt.close()

# Now update the Word document with the diagram
print("\n📝 Adding diagram to Word document...")

from docx import Document
from docx.shared import Inches

# Open existing Word document
doc = Document(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)

# Find and replace the placeholder paragraph
for i, para in enumerate(doc.paragraphs):
    if "[Note: Architecture diagram placeholder" in para.text:
        # Remove the placeholder paragraph
        p = para._element
        p.getparent().remove(p)

        # Insert the image before where the paragraph was
        # We'll add it to a new paragraph
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add image to document
        image_para = doc.add_paragraph()
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_run = image_para.add_run()
        image_run.add_picture(
            "/Users/jaysmacbook/Documents/Autumn25/capstone/architecture_diagram.png",
            width=Inches(6),
        )
        break

# Save updated document
doc.save(
    "/Users/jaysmacbook/Documents/Autumn25/capstone/Blog_Post_Event_Analytics_Pipeline.docx"
)
print("✅ Word document updated with architecture diagram")

print("\n" + "=" * 80)
print("ARCHITECTURE DIAGRAM CREATION COMPLETE")
print("=" * 80)
print("\n📊 Files created:")
print("  • architecture_diagram.png (high-resolution)")
print("  • Blog_Post_Event_Analytics_Pipeline.docx (updated with diagram)")
print("\n✅ Ready for submission!")
